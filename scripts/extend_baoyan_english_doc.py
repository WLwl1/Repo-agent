from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "Repo-Agent-保研英语面试全题库.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(92, 101, 112)
GOLD = RGBColor(122, 90, 0)


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), "Calibri")
    r_pr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_question(doc, number, en, zh, answer, follow_up=None):
    heading = doc.add_paragraph(style="Heading 3")
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run(f"Q{number}. {en}")
    set_run_font(run, size=11.5, color=DARK_BLUE, bold=True)

    cn = doc.add_paragraph(style="QuestionCN")
    cn.paragraph_format.keep_with_next = True
    run = cn.add_run(zh)
    set_run_font(run, size=9.5, color=MUTED, italic=True)

    ans = doc.add_paragraph(style="Answer")
    label = ans.add_run("Oral answer  ")
    set_run_font(label, color=BLUE, bold=True)
    run = ans.add_run(answer)
    set_run_font(run)

    if follow_up:
        tip = doc.add_paragraph(style="Tip")
        label = tip.add_run("If pressed  ")
        set_run_font(label, color=GOLD, bold=True)
        run = tip.add_run(follow_up)
        set_run_font(run, color=MUTED)


def add_script(doc, title, paragraphs):
    doc.add_heading(title, level=3)
    for i, text in enumerate(paragraphs, 1):
        p = doc.add_paragraph(style="Answer")
        label = p.add_run(f"Part {i}  ")
        set_run_font(label, color=BLUE, bold=True)
        run = p.add_run(text)
        set_run_font(run)


PITCH_AND_FLOW = [
    (
        "What is Repo Agent in one complete answer?",
        "请完整说明 Repo Agent 是什么。",
        "Repo Agent is an evidence-first repository investigation system for the step before an AI edits code. It parses a repository into symbols and typed relations, retrieves candidate files and functions from several views, uses route and graph evidence to rerank them, and returns an inspectable answer with line ranges, paths, warnings, decoys, and a replayable proof bundle. Its main purpose is to reduce code changes based on the wrong repository context.",
        "I would describe it as a localization and evidence layer, not as a complete autonomous coding agent.",
    ),
    (
        "Who are the intended users and what is the main use case?",
        "目标用户是谁？最主要的使用场景是什么？",
        "The intended users are developers, reviewers, coding agents, and researchers who need to understand an unfamiliar repository. The main use case is a question such as: which function implements this endpoint, which file should be changed, or which execution path explains this behavior? Repo Agent is most useful when the repository contains similar public, admin, legacy, mock, test, or documentation code that can mislead a simple search.",
        None,
    ),
    (
        "Can you walk through the complete request pipeline?",
        "请口述一次完整请求经过的流程。",
        "A request enters through the CLI, Web API, or MCP server. The security layer validates the repository path, question length, and top-k. The runtime loads or builds the index. Parsers extract symbols and relations; the indexer creates chunks, file facts, and a typed graph. Query planning selects an intent. File scouting and multi-view retrieval generate candidates. Route anchoring, bounded PPR, and reranking refine them. The agent then builds the answer, diagnostics, proof, bundle, or HTML report.",
        None,
    ),
    (
        "Can you explain the public /api/chat writer example?",
        "请用 public /api/chat writer 案例解释项目。",
        "The question asks which function finally writes streamed tokens for the public /api/chat endpoint. Several functions contain chat, stream, and write, including admin and legacy decoys. Repo Agent first recognizes the exact route literal. It follows the path from the route to the public handler, then to the streaming function and finally to writeChatDelta. The proof records this path and explains why writeAdminChatDelta and writeLegacyChatDelta belong to different route families.",
        "The example is strong because the correct answer is supported by an execution path, not only by shared words.",
    ),
    (
        "How do you formulate the project as a research problem?",
        "这个项目如何被建模为一个科研问题？",
        "The research problem is repository-level localization under hard negatives. Given a repository, a natural-language question, and optional route or intent constraints, the system must rank the relevant file or symbol and provide evidence for that ranking. The scientific questions are whether structured multi-view retrieval improves localization, whether replayable evidence reduces high-confidence mistakes, and whether that evidence improves downstream repair under a controlled model and token budget.",
        None,
    ),
    (
        "What does Repo Agent deliberately not do?",
        "Repo Agent 明确不做什么？",
        "It does not prove complete program semantics, guarantee that the top hit is correct, or replace full dynamic analysis. Its default mode does not train a model, require a GPU, or automatically edit the source repository. The engineering mode can support controlled changes, but only through a separate workspace and explicit apply step. I also do not claim external state of the art or improved patch success because those experiments are not yet complete.",
        None,
    ),
    (
        "Why is the evidence-first order important?",
        "为什么必须先证据、后修改？",
        "A coding model can generate a plausible patch even when it reads the wrong file. Once that happens, later reasoning is built on a false premise. The evidence-first order separates localization from modification. It gives a human or another agent a chance to inspect the target, supporting path, warnings, and alternatives before code is changed. This also makes failure analysis clearer because we can tell whether the problem came from retrieval, reasoning, or the patch itself.",
        None,
    ),
    (
        "What are the main architectural layers?",
        "项目的主要架构层有哪些？",
        "I divide the system into eight layers: security and configuration; parsing and structural extraction; data models and repository indexing; retrieval and graph ranking; agent answer construction and tools; proof, replay, and evidence artifacts; runtime and user interfaces; and optional engineering and governance extensions. This separation lets the deterministic retrieval path run without a model while higher-level interfaces reuse the same index and evidence contracts.",
        None,
    ),
    (
        "What exactly is returned to the user?",
        "系统最终向用户返回哪些内容？",
        "The public result contains the cleaned query, answer text, ranked RetrievalHit objects, file paths, symbols, line ranges, snippets, ranking reasons, a trace of the investigation stages, repository statistics, confidence diagnostics, graph-search metadata, and a proof object. Depending on the command, the same result can also be rendered as Markdown, JSON, an HTML report, an MCP payload, an impact report, or a regression contract.",
        None,
    ),
    (
        "Which components are core and which are research extensions?",
        "哪些是核心能力，哪些是研究型扩展？",
        "The core is parsing, symbol-aware chunking, multi-view retrieval, query planning, bounded graph ranking, deterministic answers, diagnostics, and proof generation. Optional interfaces include model reranking and tool calling. Impact analysis, regression contracts, temporal repair, proof attacks, the evidence court, frontier analysis, and engineering mode are extensions built on the same evidence schema. They strengthen the research story, but the project should still be explained from the core localization problem first.",
        None,
    ),
]


MODELS_CONFIG_CACHE = [
    (
        "Why do you use dataclasses instead of plain dictionaries?",
        "为什么内部数据结构使用 dataclass，而不是普通 dict？",
        "Dataclasses give the core algorithms explicit fields, type hints, defaults, and a stable place for invariants. That reduces mistakes when symbols, chunks, hits, and proof objects move across modules. At the API boundary, serializers convert them into JSON-compatible dictionaries. The cost is schema evolution: every new field may require changes to payload conversion, cache versions, MCP output, tests, and reports.",
        None,
    ),
    (
        "What information is stored in a Symbol?",
        "Symbol 数据结构保存什么信息？",
        "A Symbol represents a function, class, method, route, or another structural unit. It stores the name, kind, start and end lines, calls, references, inheritance, and an optional route path. The parser creates Symbols, and the indexer later turns them into CodeChunks. Symbol line boundaries are important because the final answer must point to a reviewable code region rather than only a file.",
        None,
    ),
    (
        "What is SourceAnalysis and why is parser_backend important?",
        "SourceAnalysis 是什么？parser_backend 为什么重要？",
        "SourceAnalysis is the per-file parser result. It contains the language, imports, extracted symbols, and parser backend. The backend may be python-ast, Tree-sitter, a regex fallback, or an error state. Keeping this provenance matters because two apparently similar chunks may have different structural reliability. It also lets evaluation separate ranking failures from parser fallback or syntax failures.",
        None,
    ),
    (
        "What makes a CodeChunk the central retrieval unit?",
        "为什么 CodeChunk 是核心检索单元？",
        "A CodeChunk connects text with structure. It contains a stable chunk id, relative path, language, code text, line range, symbol name and kind, route path, imports, calls, references, inheritance, and parser provenance. Retrieval scores the text and metadata, while the graph connects chunk ids. Because the final evidence uses the same unit, ranking, graph paths, snippets, and replay can share one identity.",
        None,
    ),
    (
        "What are FileFact and QueryPlan used for?",
        "FileFact 和 QueryPlan 分别有什么作用？",
        "FileFact summarizes a file before symbol-level retrieval. It records language, line count, imports, symbols, routes, and roles such as backend, test, config, or web. QueryPlan summarizes the question: its mode, intent, focus terms, target roles and languages, route literals, hop budget, and target symbol type. FileFact supports file scouting, while QueryPlan makes the ranking policy dependent on the actual question.",
        None,
    ),
    (
        "How do RetrievalHit, InvestigationBundle, and AgentResult differ?",
        "RetrievalHit、InvestigationBundle 和 AgentResult 有什么区别？",
        "RetrievalHit is one ranked candidate with its chunk, score, matched terms, and reasons. InvestigationBundle is the internal investigation state, including seed hits, final hits, graph edges, trace, diagnostics, and proof data. AgentResult is the public output used by CLI, Web, and MCP. The separation keeps ranking details available internally without forcing every interface to reimplement the pipeline.",
        None,
    ),
    (
        "Which configuration limits protect the system?",
        "系统有哪些关键配置上限？",
        "The current configuration limits a question to five hundred characters, top-k to twelve, indexed files to twenty-five hundred, and a source file to five hundred and twelve KiB. It also defines project, workspace, and allowed roots, plus audit and output paths. These values bound cost and attack surface. They are engineering defaults, not theoretical limits, and production deployment would need stronger configuration validation.",
        None,
    ),
    (
        "Why are ignore rules part of correctness, not only performance?",
        "为什么 ignore 规则不仅影响性能，也影响正确性？",
        "Generated reports, caches, run workspaces, dependencies, and archived artifacts often repeat the same symbols and questions as the source code. If they are indexed, they can become high-scoring documentation or duplicate decoys. Ignore rules therefore improve both speed and ranking validity. They also protect secrets such as .env files and prevent the system from investigating its own generated evidence as if it were production code.",
        None,
    ),
    (
        "How do the two cache layers differ?",
        "两层缓存有什么区别？",
        "The JSON index cache stores the complete serialized RepositoryIndex, including chunks, graph data, and retrieval state. The SQLite parse cache stores SourceAnalysis for individual files. A parse-cache hit can save AST or Tree-sitter work even when the full index must be rebuilt because graph weights or schemas changed. The two caches have separate schema versions, so their invalidation rules are intentionally different.",
        None,
    ),
    (
        "How is cache invalidation implemented and what is its limitation?",
        "缓存如何失效？当前方案有什么局限？",
        "The runtime builds a signature from supported relative paths, modification times, and file sizes after applying ignore rules. It also checks the schema version and relevant embedding configuration. This is efficient, but it is not fully content-addressed. In an extreme case, content could change while the timestamp and size remain the same. A stronger design would include per-file content hashes and support incremental edge rebuilding.",
        None,
    ),
]


PARSING_GRAPH = [
    (
        "How does the project detect supported languages?",
        "项目如何检测并支持不同语言？",
        "Language detection is mainly extension-based. The supported source types include Python, JavaScript, TypeScript, JSX or TSX variants, HTML, CSS, TOML, and selected manifest files. The parser entry point chooses the appropriate backend and records the result. Unsupported or ignored files may still appear in repository metadata, but they do not become normal source-code chunks.",
        None,
    ),
    (
        "How does the Python parser work?",
        "Python parser 的工作流程是什么？",
        "The Python path uses the standard ast module. It parses the file, records Import and ImportFrom nodes, creates symbols for functions, async functions, and classes, and walks each symbol to collect calls, references, inheritance, and route decorators. The AST provides reliable line positions and avoids matching code-like text inside comments or strings. A SyntaxError becomes an error backend with limited structure instead of crashing the entire index.",
        None,
    ),
    (
        "How are Python routes extracted?",
        "Python 路由是如何抽取的？",
        "The parser examines decorators such as route, api_route, get, post, put, patch, and delete. It extracts literal paths, HTTP methods, and the decorated handler name when those values are statically visible. The route becomes a separate Symbol and later receives a routes_to edge to the handler. Dynamic decorator construction or computed route strings may be missed, which is why route coverage must be reported rather than assumed.",
        None,
    ),
    (
        "How does the JavaScript and TypeScript parser work?",
        "JavaScript 和 TypeScript 如何解析？",
        "The preferred backend uses Tree-sitter grammars. It builds a concrete syntax tree and iteratively walks nodes to extract imports, require and dynamic import calls, functions, methods, call expressions, identifiers, inheritance, and Express-like route patterns. It also creates qualified names for class methods. Tree-sitter improves syntax tolerance, but the call graph is still an approximate name-resolution graph rather than a compiler-grade interprocedural analysis.",
        None,
    ),
    (
        "Why does the project keep a regex fallback?",
        "为什么还保留正则 fallback？",
        "Some large or template-heavy JavaScript files caused stability problems in native parsing. The project therefore uses size and safety gates and can fall back to segmented regular-expression extraction. The fallback sacrifices structural precision but prevents one file from terminating a full benchmark run. Its output is marked with parser provenance, so a successful query is not silently treated as equivalent to a full Tree-sitter analysis.",
        None,
    ),
    (
        "Why is there a large-file safety threshold?",
        "为什么对大文件设置安全阈值？",
        "The threshold limits parser memory, traversal depth, and worst-case behavior on generated or template-heavy files. In the current JavaScript path, files around or above twenty KiB may use a guarded fallback, while the general source limit is five hundred and twelve KiB. These are practical stability choices. They should be evaluated by parser recall and latency, not presented as universally optimal constants.",
        None,
    ),
    (
        "What structure is extracted from HTML, CSS, and configuration files?",
        "HTML、CSS 和配置文件会抽取哪些结构？",
        "HTML mainly contributes src and href links, CSS contributes @import relations, and TOML or manifest files are mostly represented through text and file-overview chunks. The current system does not infer a full browser event graph from HTML. Front-end behavior such as button handlers must usually be found in JavaScript. This is a clear limitation for UI-flow questions.",
        None,
    ),
    (
        "What are the three major parser failure types?",
        "解析层的三类主要错误是什么？",
        "First, syntax parsing can fail, producing an error backend. Second, parsing can succeed but structural rules may miss dynamic calls, decorators, or framework semantics. Third, symbols may be extracted correctly but graph resolution can connect a name to the wrong target. These failure types require different fixes: parser stability, extraction coverage, or edge-resolution precision.",
        None,
    ),
    (
        "Can Tree-sitter prove that the call graph is correct?",
        "Tree-sitter 能保证调用图正确吗？",
        "No. Tree-sitter provides a robust syntax tree, but Repo Agent builds calls from observed names and lightweight resolution rules. Aliases, dynamic dispatch, dependency injection, reflection, and overloaded names can still create missing or incorrect edges. I treat the graph as retrieval evidence, not ground truth, and I would evaluate edge precision and recall separately from final localization accuracy.",
        None,
    ),
    (
        "How are duplicate symbol names resolved?",
        "同名函数或方法如何解析？",
        "Resolution first prefers a qualified name, then a lowercase symbol name. If several candidates share the name, a same-file candidate is preferred; otherwise the system uses available global candidates and file or import context. This is transparent but heuristic. Cross-package duplicates and methods with the same short name remain difficult, so ranking should preserve ambiguity and avoid pretending that one edge is certain.",
        None,
    ),
    (
        "Which graph edge types are used?",
        "仓库图包含哪些边？",
        "The main edge types are calls, references, inherits, imports, and routes_to. Same-file call edges currently receive stronger weight than cross-file calls; inheritance and route-to-handler edges are also strong signals. References are weaker because name occurrence does not necessarily imply execution. Edge labels and weights are preserved so the system can explain why a graph path contributed to a result.",
        None,
    ),
    (
        "How are import edges resolved and where can they fail?",
        "import 边如何解析？哪些情况会失败？",
        "The lightweight resolver often compares the last component of an imported module with file stems or known paths. This works for many local modules but can fail with aliases, re-exports, namespace packages, generated modules, or two files sharing the same stem. A stronger implementation would use language-specific module resolution, environment metadata, and edge provenance.",
        None,
    ),
    (
        "Why do you use symbol-based chunks instead of fixed token windows?",
        "为什么按 symbol 切块，而不是固定 token？",
        "A fixed window can split one function across chunks or mix several functions in the same chunk. Symbol chunks preserve a reviewable unit, line range, calls, references, and route metadata. That improves localization and evidence handoff. The trade-off is that very long functions remain long, nested symbols may overlap, and parser failures can remove the ideal boundary. File-overview chunks provide a fallback.",
        None,
    ),
    (
        "Why is there also a file-overview chunk?",
        "既然按 symbol 切块，为什么还需要 file overview？",
        "A file overview preserves module-level imports, constants, top-level context, and files with few recognized symbols. It also protects recall when the parser misses a function. The current overview uses up to the first one hundred and forty lines. Because it can duplicate symbol text, the reranker downweights overview chunks for function or flow questions and limits their influence.",
        None,
    ),
    (
        "How do file-size and file-count limits affect recall?",
        "文件大小和数量上限如何影响召回？",
        "The limits prevent unbounded indexing, but they can exclude a true target in a very large repository. The system should report skipped files and coverage instead of treating the index as complete. For monorepos, I would replace a hard global limit with incremental indexing, directory or ownership partitions, and an explicit scout phase that can request additional regions when evidence is weak.",
        None,
    ),
    (
        "How would you evaluate parser and graph quality directly?",
        "如何单独评测 parser 和图的质量？",
        "I would sample repositories by language and framework, create gold symbols and edges, and report symbol recall, route recall, call-edge precision and recall, and fallback rates. I would also stratify localization results by parser backend. This separates a retrieval failure from a missing representation and prevents the final Top-k metric from hiding parser weaknesses.",
        None,
    ),
]


RETRIEVAL_RANKING = [
    (
        "Can you explain the BM25 formula orally?",
        "请口述 BM25 公式的含义。",
        "BM25 gives a document more credit when it contains a rare query term, but the benefit of repeating the term gradually saturates. It also normalizes for document length, so a long function does not win only because it repeats common words. In the implementation, inverse document frequency uses the number of documents containing the term, and the term-frequency fraction is controlled by k1 and b.",
        None,
    ),
    (
        "What do k1 and b control?",
        "BM25 的 k1 和 b 分别控制什么？",
        "k1 controls how quickly term-frequency gain saturates. A larger value allows repeated terms to keep contributing for longer. b controls document-length normalization: zero ignores length, while one applies full normalization relative to the average document length. Repo Agent currently uses k1 equal to 1.5 and b equal to 0.75, common starting values that still require held-out evaluation.",
        None,
    ),
    (
        "How are queries tokenized and expanded?",
        "查询如何分词和扩展？",
        "The deterministic path normalizes terms, extracts identifiers and route literals, and expands selected intent words with a small transparent vocabulary. For example, a streaming-writer question may add related action terms, while a configuration question emphasizes package and metadata roles. Expansion improves recall but creates an overfitting risk. Every rule should be versioned, logged, ablated, and tuned only on development data.",
        None,
    ),
    (
        "Why are the four retrieval views independent?",
        "为什么四个检索视图必须独立？",
        "Content, identifiers, paths, and structure express different evidence. If they are concatenated, a long function body can dominate a short exact symbol or path, and it becomes difficult to measure which evidence helped. Independent indexes preserve separate rankings, allow per-view ablation, and let rank fusion reward agreement without forcing the raw scores onto one scale.",
        None,
    ),
    (
        "Can you explain weighted reciprocal rank fusion?",
        "请解释 weighted RRF。",
        "Each view first returns an independent ranking. For a candidate at rank r in a view, RRF adds the view weight divided by a constant plus r. Repo Agent uses a rank constant of thirty and current weights of 1.0 for content, 1.8 for identifiers, 1.1 for paths, and 1.25 for structure. The method is stable and interpretable, but it discards raw-score magnitude.",
        None,
    ),
    (
        "Why is the RRF rank constant set to thirty?",
        "RRF 的 rank constant 为什么设为 30？",
        "The constant controls how quickly rank contributions decay. A smaller value emphasizes the first few positions more strongly; a larger value makes lower ranks more similar. Thirty is an engineering choice that worked in the current development setting, not a theoretical optimum. It should be frozen after development tuning and tested through sensitivity analysis on repository-disjoint data.",
        None,
    ),
    (
        "How were the four view weights selected?",
        "四个 view 的权重如何确定？",
        "The current weights reflect the idea that exact identifiers are especially valuable in code, while content, structure, and path provide complementary evidence. They were refined through project-controlled cases. That creates a possible tuning bias. A stronger study would report single-view and pairwise ablations, sensitivity curves, rule counts, and performance on a frozen external test before claiming the weights generalize.",
        None,
    ),
    (
        "What does DenseEmbeddingIndex do?",
        "DenseEmbeddingIndex 的作用是什么？",
        "It is an optional in-memory cosine-similarity index. It validates vector dimensions, stores vector norms, and ranks chunk vectors against a query vector. It can recover semantic matches that do not share exact words. However, it is not an approximate-nearest-neighbor database and does not make the default system dependent on a model. Dense results are an additional channel, not a replacement for exact code evidence.",
        None,
    ),
    (
        "Why keep lexical retrieval when dense retrieval is available?",
        "有向量检索时为什么仍然保留词法检索？",
        "Code contains exact signals such as function names, paths, package keys, and route literals. Dense similarity can blur those distinctions or rank a semantically similar but operationally wrong function. Lexical retrieval is cheap, local, deterministic, and easy to audit. The best design is usually hybrid: use dense retrieval for vocabulary mismatch and lexical plus structural evidence for precise grounding.",
        None,
    ),
    (
        "What intents can the query planner recognize?",
        "Query planner 能识别哪些意图？",
        "Representative intents include front-end lookup, style lookup, test lookup, configuration lookup, API lookup, flow tracing, and general code search. Each intent suggests target file roles, languages, symbol kinds, action words, and graph depth. The planner is rule-based, so it is deterministic and debuggable, but it may miss paraphrases or encode project-specific vocabulary.",
        None,
    ),
    (
        "Why use both file scouting and global chunk recall?",
        "为什么同时需要 file scout 和全局 chunk recall？",
        "File scouting cheaply identifies likely modules and roles, which reduces noise. But a hard file-level filter can permanently remove the correct symbol. Repo Agent therefore combines chunks from scouted files with the top globally scored chunks. This union separates file-level recall from symbol-level precision and gives the second stage a recovery path when the first stage is wrong.",
        None,
    ),
    (
        "How are candidate budgets bounded?",
        "候选数量如何限制？",
        "The scout expands a bounded number of files based on top-k, while global scoring contributes a bounded set of chunks. Graph expansion also limits the number of edges and local nodes. These budgets keep latency predictable and prevent high-degree modules from consuming the search. The exact constants are engineering choices and should be reported with latency, memory, and accuracy sensitivity.",
        None,
    ),
    (
        "How does route anchoring influence ranking?",
        "route anchor 如何影响排序？",
        "An exact route literal identifies a route node and route family. Candidates reachable from that route through handler and call edges receive evidence support. Candidates associated with different route families, such as admin or legacy paths, can be penalized even if their words are similar. Route anchoring is an independent signal, so it is not completely replaced by a high lexical score.",
        None,
    ),
    (
        "What is contrastive reranking?",
        "什么是 contrastive reranking？",
        "Contrastive reranking asks not only why a candidate matches, but also why a tempting alternative conflicts with the query. Terms such as public, not admin, current, or non-legacy can define exclusions. File roles, route families, and decoy labels then reduce the score of incompatible candidates. This is useful for hard negatives, but the rules must be tested against cases where an excluded word appears in valid context.",
        None,
    ),
    (
        "What does multistep reranking add after retrieval?",
        "multistep rerank 在初始检索之后增加了什么？",
        "It combines file-role alignment, language alignment, symbol-kind targets, action matches such as write or stream, route reachability, graph evidence, overview penalties, test or documentation penalties, and contrastive exclusions. The purpose is to convert general textual relevance into question-specific relevance. Each reason is attached to the hit so the score is not an unexplained black box.",
        None,
    ),
    (
        "How does Personalized PageRank work in this project?",
        "PPR 在项目中的数学和工程流程是什么？",
        "Initial retrieval scores define the teleport distribution. At each iteration, a fraction of relevance returns to those seeds, while the rest moves through normalized weighted graph edges. The damping factor is 0.85. The current implementation uses a bounded neighborhood, at most eighty iterations, and an L1 convergence threshold of ten to the minus seven. The resulting relevance becomes a controlled rerank boost.",
        None,
    ),
    (
        "How do you prevent graph evidence from overwhelming lexical evidence?",
        "如何防止 graph boost 淹没词法证据？",
        "The graph operates on candidates seeded by retrieval, uses a bounded neighborhood, normalizes probability mass, and contributes a limited boost rather than replacing the base score. Route anchors and intent rules are also evaluated separately. I would still verify this through graph-off ablation and score-distribution inspection because a badly connected utility module can otherwise accumulate relevance.",
        None,
    ),
    (
        "Why does the code still contain MCTS names?",
        "为什么代码和 artifact 中仍然有 MCTS 命名？",
        "The project previously experimented with a pseudo-MCTS or greedy graph walk, and public schemas kept names such as graph_mcts for compatibility. The active algorithm is bounded PPR. This is technical debt, not a reason to relabel PPR as MCTS. A proper cleanup needs a schema version, migration path, backward-compatible reader, and updated tests and reports.",
        None,
    ),
]


AGENT_RUNTIME_INTERFACES = [
    (
        "What are the two answer paths in RepoAgent.answer?",
        "RepoAgent.answer 有哪两条回答路径？",
        "The default path is deterministic: it uses the investigation bundle to generate a grounded answer without a network model. The optional path can use model reranking or a bounded tool-calling loop when a compatible model is configured. Both paths must remain grounded in retrieved or observed repository evidence. The model is not allowed to invent a symbol that was never found or read.",
        None,
    ),
    (
        "Why keep a deterministic answer mode?",
        "为什么保留 deterministic answer？",
        "It provides a local, reproducible, low-cost baseline and keeps the core system usable without an API key. It also makes evaluation easier because ranking changes are not mixed with model randomness. The limitation is that the prose can be less flexible and may not synthesize complex context as well as a model. I treat it as a reliable baseline, not as the only user experience.",
        None,
    ),
    (
        "What is the optional LLM reranker allowed to do?",
        "可选 LLM reranker 可以做什么？",
        "It receives an already retrieved candidate set and can compare those candidates against the question, similar to a cross-encoder. It should not perform unbounded repository discovery or introduce unseen symbols. This boundary keeps recall attributable to the deterministic pipeline and makes model value measurable as a reranking layer.",
        None,
    ),
    (
        "Where is the agent behavior in this project?",
        "这个项目的 Agent 性体现在哪里？",
        "When model mode is enabled, the system can plan, call bounded repository tools, observe results, choose follow-up reads or relation searches, and stop with an evidence-backed answer. Even without a model, the pipeline has staged planning and verification. I avoid calling every function an agent; the strongest agent behavior is the explicit tool-observation loop and the engineering run state machine.",
        None,
    ),
    (
        "Why not send the whole repository in one large prompt?",
        "为什么不把整个仓库塞进一个大 prompt？",
        "Large prompts are expensive, may exceed context limits, and dilute exact evidence with irrelevant files. They also make it difficult to explain why a particular function mattered or to replay the decision after code changes. Retrieval and graph evidence reduce the context to a reviewable set. A model can then reason over that set with clearer provenance and a controlled token budget.",
        None,
    ),
    (
        "Which read-only tools are available to the agent?",
        "Agent 有哪些只读工具？",
        "Representative tools include repository memory and brief, query planning, semantic scoring, file scouting, candidate reading, graph-neighbor following, reranking, relevant-edge extraction, directory listing, text search, symbol search, relation search, and file reading. Tool outputs are recorded in the trace. The interface gives the model structured observations instead of unrestricted shell access.",
        None,
    ),
    (
        "Which tools can write or execute commands?",
        "哪些工具可以写文件或执行命令？",
        "The sensitive tools are replace_text, write_file, and run_command. They must resolve paths inside the allowed repository or workspace, reject ignored and protected locations, and apply the command allowlist. In the normal investigation mode, the project emphasizes read-only evidence. Writing belongs to the separate engineering workflow with state, review, and apply controls.",
        None,
    ),
    (
        "What are repository memory and repository brief?",
        "repo memory 和 repo brief 是什么？",
        "They are compact summaries that help an agent avoid rediscovering basic repository facts in every turn. The brief can include languages, likely entry points, roles, and index statistics. Memory can preserve bounded observations from the current investigation. They are not a substitute for reading evidence, and stale summaries must be tied to the current index revision.",
        None,
    ),
    (
        "What does RepoAgentRuntime centralize?",
        "RepoAgentRuntime 统一管理什么？",
        "It centralizes configuration, allowed roots, reports and runs directories, cache objects, audit logging, optional model clients, in-memory indexes, and the public operations for load, ask, report, bundle, impact, engineering, and apply. CLI, Web, and MCP reuse this lifecycle so they do not develop different security or indexing semantics.",
        None,
    ),
    (
        "What is the exact load_index order?",
        "load_index 的执行顺序是什么？",
        "It validates the path, computes the repository signature, checks the in-memory cache, then checks the disk JSON index cache and schema. If a rebuild is required, build_index scans supported files, reuses per-file parse-cache entries where valid, constructs chunks and graph edges, builds retrieval indexes, writes the new cache, and records an audit event. Model or embedding configuration changes can also invalidate reuse.",
        None,
    ),
    (
        "What is the difference between a bundle and a report?",
        "bundle 与 report 有什么区别？",
        "A bundle is the evidence handoff artifact, usually JSON or Markdown, designed for another tool, agent, or reviewer. A report is a human-oriented HTML presentation with panels and graphs. The bundle is closer to the source of truth. If a visual report looks wrong, I first inspect the bundle and serializer before blaming retrieval.",
        None,
    ),
    (
        "What should an HTML report show?",
        "HTML 报告应展示哪些内容？",
        "It should show the question, top hits, snippets and reasons, graph-search audit, proof status and supporting paths, confidence warnings, proof graph, and contrastive decoy audit. The goal is not decoration; it is to make the evidence reviewable. HTML rendering must escape untrusted text and preserve the index revision or fingerprint used to generate the report.",
        None,
    ),
    (
        "What is recorded by the audit logger?",
        "审计日志记录什么？",
        "The audit layer records important runtime events such as path validation, cache loads or rebuilds, asks, tool use, report generation, engineering state transitions, and apply operations. It supports debugging and accountability. Audit logs can contain repository paths or prompts, so retention, access control, and privacy must be considered in a production deployment.",
        None,
    ),
    (
        "What happens when no model configuration is available?",
        "没有模型配置时会发生什么？",
        "The core index, retrieval, graph ranking, deterministic answer, proof, replay, reports, and benchmarks still work. Model-specific paths return a clear unavailable or fallback trace rather than pretending a model ran. Engineering mode may only plan or use limited fallback behavior. A successful JSON response does not prove that an autonomous code change was completed.",
        None,
    ),
    (
        "Is the shared runtime safe under concurrent requests?",
        "共享 runtime 在并发请求下安全吗？",
        "The runtime reuses in-memory indexes and shared directories, so concurrency requires careful locking and immutable index objects. The current local prototype is not presented as a hardened multi-tenant service. A production version would add request isolation, cache synchronization, bounded worker pools, unique run directories, authentication, and tests for concurrent rebuild and report generation.",
        None,
    ),
    (
        "How do you diagnose an incorrect answer?",
        "系统答错时按什么顺序诊断？",
        "I check the query plan first, then whether the correct file was indexed, whether the parser created the symbol and edges, whether file scouting or global recall included it, how each retrieval view ranked it, what PPR and reranking changed, and finally how the answer and report were serialized. This order separates representation, retrieval, graph, policy, and presentation failures.",
        None,
    ),
]


PROOF_GOVERNANCE = [
    (
        "What fields are stored in a proof bundle?",
        "proof bundle 保存哪些字段？",
        "It stores the query and plan, seed and final hits, the top hit, route anchors, supporting paths, proof-graph nodes and edges, snippets or fingerprints, decoy audit, score gaps, warnings, and status checks. These fields allow a later replay to verify the evidence structure without repeating the original natural-language reasoning.",
        None,
    ),
    (
        "What does proof status actually assert?",
        "proof status 实际断言了什么？",
        "A proved status means that the required evidence objects in the current schema resolve and pass their checks, such as the top hit, route anchor, supporting path, and selected graph edges. It does not assert complete semantic correctness, runtime reachability under every condition, or patch safety. I sometimes describe it as an evidence contract to avoid overclaiming.",
        None,
    ),
    (
        "What levels of proof replay are supported?",
        "proof replay 有哪些检查层级？",
        "Basic replay checks whether saved targets, routes, snippets, and supporting nodes still exist. Strict replay also verifies that saved route and path edges correspond to current graph edges. Scorecards summarize several replay and mutation checks. The levels are useful because a node can survive a refactor while the evidence path connecting it to the original behavior has broken.",
        None,
    ),
    (
        "How is proof drift diagnosed?",
        "证据漂移如何分类？",
        "The system can distinguish top-hit drift, route-anchor drift, execution-path drift, stale proof-graph endpoints, and decoy-audit drift. Each category suggests a different response: rerun localization, repair the parser or graph, inspect a moved symbol, or update the decoy set. Drift diagnosis is more useful than returning only pass or fail.",
        None,
    ),
    (
        "What does the mutation lab test?",
        "mutation lab 测试什么？",
        "It creates controlled corruptions such as a stale top hit, missing route, broken supporting path, invalid proof edge, or changed decoy entry, then checks whether replay detects them. A high mutation-detection rate shows that the verifier is not designed to always pass. It still does not prove coverage of every real repository change.",
        None,
    ),
    (
        "What are the limitations of proof-carrying retrieval?",
        "Proof-Carrying Retrieval 有哪些边界？",
        "The proof depends on the quality of parsing, graph construction, labels, and the selected evidence schema. It can be self-consistent while the repository model is incomplete. Hashes detect change but do not establish trust in the original artifact. Replay also cannot replace tests or dynamic analysis. Its value is traceability and drift detection, not mathematical program verification.",
        None,
    ),
    (
        "How is a proof graph different from the repository graph?",
        "proof graph 与 repository graph 有什么区别？",
        "The repository graph contains all indexed chunks and relations. The proof graph is a small query-specific subgraph containing the route anchor, top hit, supporting path, visited or relevant nodes, and selected decoys. It is designed for explanation and replay. Because it is a projection, omitting a node does not mean that node is absent from the repository graph.",
        None,
    ),
    (
        "Why is a decoy audit necessary?",
        "为什么需要 decoy audit？",
        "A top result can look convincing while a nearly identical wrong implementation remains unexamined. The decoy audit records those alternatives, their ranks, route families, and rejection signals. This supports red-team evaluation and helps a reviewer see whether the system understood the distinction. The current audit coverage is not complete, so missing decoys should remain a warning.",
        None,
    ),
    (
        "Is the confidence score a calibrated probability?",
        "confidence score 是校准后的概率吗？",
        "No. It is a diagnostic heuristic based on evidence count, score gap, graph and route support, proof status, and warnings. A value such as 0.92 should not automatically be interpreted as a 92 percent correctness probability. Calibration requires held-out predictions, reliability diagrams, Brier or calibration error, and risk-coverage evaluation.",
        None,
    ),
    (
        "How does impact analysis use a proof bundle?",
        "impact analysis 如何利用 proof bundle？",
        "It starts from the proved target and walks upstream and downstream graph relations to identify exposed routes, related files, dependents, risk items, and a verification plan. This connects localization with change reasoning. Because the graph is approximate, the report should be treated as a prioritized review surface, not a complete dependency analysis.",
        None,
    ),
    (
        "What is a proof regression contract?",
        "什么是 proof regression contract？",
        "It freezes important evidence invariants such as the target symbol, route literals, supporting path, graph edges, decoy expectations, and exposed surfaces. A later verification checks whether a change preserved those invariants. The contract turns a one-time investigation into a reviewable PR-time guard, but it must be migrated when a legitimate refactor changes the evidence structure.",
        None,
    ),
    (
        "What does the PR guard do?",
        "PR guard 的作用是什么？",
        "Given a contract and changed files, the PR guard determines whether protected evidence surfaces were touched and which replay or verification commands are required. It can emit pass, warning, or failure decisions, GitHub annotations, and SARIF for code scanning. It does not merge code or prove behavioral correctness; it enforces that relevant evidence checks are not silently skipped.",
        None,
    ),
    (
        "What is temporal proof regression?",
        "什么是 temporal proof regression？",
        "It replays a proof contract across a range of git commits to find the last passing and first failing revision. Instead of asking only whether evidence is broken now, it asks when the break occurred. The implementation exports commit snapshots, builds the necessary repository view, and records the transition without checking out or mutating the active workspace.",
        None,
    ),
    (
        "Why export git snapshots instead of switching the worktree?",
        "为什么导出 git snapshot，而不是切换当前工作树？",
        "Using git archive or isolated snapshots avoids changing the developer's active files, branch, uncommitted work, or index. Each revision can be analyzed in a separate directory with a clear commit identity. This improves safety and reproducibility. The trade-off is additional disk and indexing cost across many commits.",
        None,
    ),
    (
        "How does successor inference work after a symbol moves?",
        "符号移动或改名后，successor inference 如何工作？",
        "It ranks possible successor symbols using route reachability, continuity from proof-path predecessors, body-token overlap, call overlap, file and name similarity, and whether the new node explains the broken graph edge. The result is a review candidate, not an automatic rewrite. Negative-control cases should force the system to abstain when no valid successor exists.",
        None,
    ),
    (
        "What is a proof graph delta?",
        "什么是 proof graph delta？",
        "It compares the evidence path in the last passing snapshot with the first failing snapshot. It reports which nodes and edges were preserved, removed, unresolved, or relinked to a possible successor. This provides a causal explanation for the regression. A successor label without a graph delta would be much weaker evidence.",
        None,
    ),
    (
        "Why use JSON Patch for contract migration?",
        "为什么用 JSON Patch 表达 contract migration？",
        "JSON Patch represents explicit add, remove, and replace operations over the contract. It makes the proposed migration reviewable and machine-applicable without silently rewriting the whole artifact. The plan can replace the target, proof context, and supporting path, then simulate whether the successor exists and the path reconnects. Human approval is still required.",
        None,
    ),
    (
        "What is the multi-agent evidence court?",
        "多 Agent 证据法庭是什么？",
        "It is a structured reliability protocol, not a group chat. Specialized roles such as retrieval advocate, graph navigator, proof verifier, mutation skeptic, red-team skeptic, temporal guardian, and arbiter emit claims with evidence hashes. Challenges are raised against unsupported or conflicting claims, and the arbiter accepts the result only when required claims pass and error-level challenges are discharged.",
        None,
    ),
    (
        "How does the arbiter reach a verdict?",
        "arbiter 如何给出 verdict？",
        "The arbiter reads the claim ledger, required role outputs, evidence hashes, challenge severity, and discharge status. It does not average agent opinions. A confident retrieval claim can be rejected if strict replay fails or an unmitigated decoy remains. This design tries to make disagreement machine-checkable, although the role policies and thresholds still require evaluation.",
        None,
    ),
    (
        "What do frontier, ablation, interaction, and stability analyses add?",
        "frontier、ablation、interaction 和 stability 分析分别增加什么？",
        "The frontier maps evidence artifacts into multiple dimensions such as reliability, robustness, adaptivity, governance, and efficiency, then identifies Pareto-efficient profiles. Single-family ablation asks which evidence family supports the result. Pairwise interaction tests whether two families fail nonlinearly together. Stability analysis perturbs metrics to estimate whether frontier membership and interactions survive uncertainty. These are diagnostic analyses, not a substitute for external task success.",
        None,
    ),
]


ENGINEERING_SECURITY = [
    (
        "What is Engineering Mode designed to do?",
        "Engineering Mode 的目标是什么？",
        "It extends investigation into a controlled change workflow. The system can plan a task, inspect evidence, edit a separate workspace, run an allowed verification command, review the diff, and preserve a structured run record. It is experimental and optional. The original repository is not modified by default.",
        None,
    ),
    (
        "Why is workspace mode the default?",
        "为什么默认在 workspace 中修改？",
        "A copied or isolated workspace protects the user's source tree and uncommitted work. It allows edits, tests, and review to happen before any apply step. It also gives every run a stable directory and record for resume or audit. The cost is copy time, storage, and the need to detect source drift before applying changes back.",
        None,
    ),
    (
        "What states does an EngineeringRun have?",
        "EngineeringRun 有哪些状态？",
        "A run records its task, workspace, events, steps, changes, verification results, review findings, and completion or failure status. Typical transitions include created, planned, investigating, editing, verifying, reviewing, completed, failed, or awaiting apply. Persisting the state allows resume and makes the agent's actions auditable instead of disappearing into chat history.",
        None,
    ),
    (
        "How do the verifier and reviewer differ?",
        "Verifier 与 Reviewer 有什么区别？",
        "The verifier selects or runs an allowed command and classifies objective command results. The reviewer examines the changed files, risk surface, missing tests, and whether the task was actually satisfied. A passing test command does not guarantee a good patch, while a review comment is not a substitute for execution. Keeping the roles separate makes their evidence clearer.",
        None,
    ),
    (
        "What makes apply-run dangerous?",
        "apply-run 的主要风险是什么？",
        "The source repository may have changed after the workspace was created, so copying files back can overwrite newer work or create semantic conflicts. A safe apply needs explicit confirmation, a known changed-file set, source-base hashes, protected-path checks, and ideally a three-way merge or patch review. The current prototype does not justify unattended apply in a high-risk repository.",
        None,
    ),
    (
        "What is the local threat model?",
        "本地 Agent 的威胁模型是什么？",
        "Threats include path traversal, symlink escape, reading secrets, indexing generated bait, unsafe command arguments, prompt injection in repository text, overwriting protected files, malicious reports, stale caches, and accidental application to the wrong repository. Running locally reduces data exposure but does not remove these risks. The system therefore centralizes path and command policy.",
        None,
    ),
    (
        "How are repository paths validated?",
        "仓库路径如何校验？",
        "The path is resolved to an absolute directory and checked against configured allowed roots. Internal file operations use safe joins and repository-relative paths. A production design must also handle symlinks, junctions, case-insensitive paths, and race conditions between validation and use. String-prefix checks alone are not sufficient.",
        None,
    ),
    (
        "Why validate question length and clamp top-k?",
        "为什么限制问题长度并 clamp top-k？",
        "These limits bound parsing, retrieval, response size, and model cost. They also reduce denial-of-service and prompt-abuse surface. The current defaults are five hundred characters and a maximum top-k of twelve. The system rejects an empty or oversized question and clamps or validates top-k rather than allocating unbounded work.",
        None,
    ),
    (
        "Why is shell=False not a complete security solution?",
        "为什么 shell=False 仍不能保证命令安全？",
        "shell=False prevents shell metacharacters from being interpreted by an intermediate shell, but the executable itself can still be dangerous. A compiler, test runner, or package manager may execute project code, load configuration, or modify files. Argument injection and executable-path spoofing are also possible. The project therefore combines shell=False with parsed arguments, an allowlist, path checks, and bounded command shapes.",
        None,
    ),
    (
        "What are the limitations of the command allowlist?",
        "命令 allowlist 有什么局限？",
        "An allowed command is not automatically safe for every repository. For example, a test command can execute malicious tests, and a Python executable can run arbitrary code if arguments are too broad. The allowlist must validate the resolved executable, subcommand, arguments, working directory, and expected side effects. Isolation and least privilege remain necessary for untrusted repositories.",
        None,
    ),
    (
        "How are protected files such as .env handled?",
        "如何处理 .env 等敏感或受保护文件？",
        "Secret and environment files should be ignored during indexing and rejected by write tools. Generated report, cache, run, and dependency directories are also protected from normal source edits. Ignore rules reduce accidental exposure, but secrets can still appear in source or logs. A production system would add secret scanning, log redaction, and stricter output retention.",
        None,
    ),
    (
        "How do you handle prompt injection inside a repository?",
        "如何防御仓库文本中的 prompt injection？",
        "Repository text is treated as untrusted data, not as an instruction. Tool permissions are enforced outside the model, and evidence snippets cannot grant new capabilities. The model receives bounded tools and observations, while path and command checks remain deterministic. Prompt injection can still bias reasoning, so suspicious instructions should be surfaced as content and excluded from control flow.",
        None,
    ),
    (
        "Why is __main__.py considered technical debt?",
        "为什么 __main__.py 被认为是技术债？",
        "It contains a very large number of CLI subcommands, renderers, benchmark orchestration, release logic, git fixtures, and compatibility paths. This makes change review, ownership, and testing harder. The core algorithms live in imported modules, so I would split the CLI into command packages with shared argument and output contracts while preserving backward compatibility.",
        None,
    ),
    (
        "What are the responsibilities of the Web Studio?",
        "Web Studio 的职责边界是什么？",
        "The server exposes local APIs for indexing, asking, reports, runs, and static assets. The front end collects the repository path and question, calls those APIs, and renders hits, traces, diagnostics, graph evidence, and proof panels. It should not reimplement ranking logic. The current Studio is a local review interface, not a hardened public internet service.",
        None,
    ),
    (
        "What does the MCP server provide?",
        "MCP server 提供什么能力？",
        "It exposes repository investigation as structured tools for an MCP-compatible host. A call includes the repository path, question, and top-k, then returns index statistics, answer, hits, trace, diagnostics, graph search, and proof data. MCP is a protocol boundary, not a model. Field names and schema versions must remain synchronized with runtime and tests.",
        None,
    ),
    (
        "How do CLI, Web, and MCP share the same semantics?",
        "CLI、Web 和 MCP 如何保证语义一致？",
        "All three interfaces call RepoAgentRuntime and the same serializers rather than building independent pipelines. Security validation, cache loading, indexing, asking, proof generation, and result fields are therefore centralized. Contract tests should compare representative outputs across interfaces. The current MCP backend-field drift is an example of why these schema tests matter.",
        None,
    ),
]


EVALUATION_RESEARCH = [
    (
        "Why does the project use several benchmark layers?",
        "为什么项目使用多层 benchmark？",
        "Each layer answers a different question. The bundled suite checks regression and known mechanisms. The challenge suite adds harder intents and distractors. Counterfactual and proof-attack suites test specific failure mechanisms. Repository-disjoint external data tests cross-repository generalization. Temporal benchmarks test drift and successor inference. Combining them into one score would hide which claim is actually supported.",
        None,
    ),
    (
        "What does the ten-case portable suite prove and not prove?",
        "10 题 portable suite 能说明什么，不能说明什么？",
        "It shows that the current pipeline solves ten controlled cases across bundled fixture repositories and remains stable as a regression gate. The recorded result is 100 percent Top-1 and Top-3 with an MRR of 1.000. Because the cases are shipped with the project and may influence rules, the suite does not establish broad external validity or state-of-the-art performance.",
        None,
    ),
    (
        "What does the thirty-two-case challenge reveal?",
        "32 题 challenge 暴露了哪些问题？",
        "It reaches 84.375 percent Top-1, 93.75 percent Top-3, and 0.880 MRR with zero distractor-at-one. The remaining errors include configuration and package-data cases, Web run-history refresh, authorization middleware, streaming-turn builders, and state-clear helpers. The result shows strong route-oriented behavior but weaker generalization to configuration, front-end state, and some cross-file intents.",
        None,
    ),
    (
        "How was the external repository-disjoint evaluation built?",
        "外部 repository-disjoint 评测如何构建？",
        "A CORE-Bench-derived collection contains two hundred queries from twenty-two repositories. Repositories, not individual queries, define the split, producing 122 train, 28 development, and 50 frozen test queries. This reduces leakage from repository-specific naming. The adapter stores provenance, expected files or symbols, tags, and hashes, but the external data and labels still require license and gold-quality auditing.",
        None,
    ),
    (
        "How should the external results be interpreted?",
        "外部实验结果应该如何解读？",
        "On the fifty-query frozen test, Full Multiview reaches 16 percent Hit@1, 32 percent Hit@3, 36 percent Hit@5, and 0.254 MRR. Single-view BM25 reaches 14, 20, 24 percent and 0.196 MRR. The structured method improves candidate placement, especially Hit@3 and Hit@5, but rank-one improvement is small and the current confidence intervals cross zero.",
        None,
    ),
    (
        "Why report Hit@k, MRR, nDCG, recall, and distractor-at-one?",
        "为什么需要同时报告多种指标？",
        "Hit@1 measures the immediate answer. Hit@3 and Hit@5 measure a short human-review list. MRR reflects the first relevant rank. nDCG supports graded relevance, and Recall@100 checks whether the correct target entered the candidate pool. Distractor-at-one measures whether a designed hard negative defeats the system. Latency and memory are also necessary for a system claim.",
        None,
    ),
    (
        "What is the baseline matrix and how do you keep it fair?",
        "baseline 矩阵是什么？如何保证公平？",
        "The internal matrix includes content-only BM25, content plus identifiers, content plus structure, Full Multiview, graph-off and graph-on variants, and optional dense or model reranking. Fair comparison requires the same repositories, queries, relevance judgments, top-k, hardware, and evaluation code, changing only the intended component. External systems also need matched context and cost budgets.",
        None,
    ),
    (
        "How do repository-disjoint splits and frozen hashes prevent leakage?",
        "repository-disjoint split 和 frozen hash 如何防止泄漏？",
        "All queries from one repository belong to one split, so repository-specific identifiers cannot appear in both tuning and testing. The frozen partition is serialized and hashed, and a verification step detects later changes. A tuning log records rules and their evidence. These controls reduce leakage, but they do not fix biased label construction or repeated upstream code across repositories.",
        None,
    ),
    (
        "What does paired bootstrap tell you?",
        "paired bootstrap 告诉了你什么？",
        "It repeatedly resamples the same query pairs and recomputes the metric difference between two methods. This estimates uncertainty without assuming normality and preserves the paired design. In the current external study, intervals for several key metrics cross zero, so the observed improvement should be described as a trend or effect estimate, not a statistically significant result.",
        None,
    ),
    (
        "Which system metrics are still required?",
        "除了准确率，还必须报告哪些系统指标？",
        "I should report index time, query latency, peak memory, index size, cache hit rate, parser fallback rate, graph size, and optional model cost and tokens. Results should be stratified by repository size, language, parser backend, and query type. Without these metrics, a ranking improvement may hide an impractical system cost.",
        None,
    ),
    (
        "How do you interpret the current proof-attack results?",
        "当前 proof-attack 结果如何解释？",
        "The three synthetic attacks currently show 100 percent attack resistance, Top-1, proof status, route-anchor preservation, and supporting-path preservation. Generated-decoy audit coverage is about 66.67 percent, and the mitigation signal rate is zero. This means the ranking resisted the examples, but the audit did not establish a clear causal defense signal for every decoy. I must report both sides.",
        None,
    ),
    (
        "What do benchmark diagnostics and repair artifacts do?",
        "benchmark diagnostics 与 repair artifacts 有什么作用？",
        "Diagnostics classify weak cases such as top-three recoverable, library-boundary ambiguity, streaming-handler ambiguity, hidden symbols, and weak route anchors. Repair synthesis converts those traces into explicit rule candidates with coverage, risk cases, and projected effects. Implementation verification checks that the source actually contains the guard and reason. Compilation and workbench artifacts create reviewable interventions and ablation plans.",
        None,
    ),
    (
        "What is the proof-attack CEGAR or minimax loop?",
        "proof-attack 的 CEGAR/minimax 闭环是什么？",
        "The system generates counterexamples, ranks attack pressure, creates defense triage, synthesizes a policy, generates adaptive attacks against that policy, proposes repairs, and records a certificate. Conceptually, the defense becomes the next attack surface. In the current workspace, several state-machine and scorecard tests fail, so this loop is a research prototype and cannot be presented as a validated security guarantee.",
        None,
    ),
    (
        "How does the release pack support reproducibility?",
        "release pack 如何支持可复现性？",
        "It collects benchmark reports, proof artifacts, scorecards, contracts, and evaluation cards, then records file paths, byte sizes, and SHA-256 hashes in a manifest. A verifier checks for missing or modified artifacts. The hash protects integrity after generation, but it does not prove the artifact was correct or produced in a trusted environment. Clean-clone reproduction and CI attestation are stronger next steps.",
        None,
    ),
    (
        "How do you explain the current full-test failures?",
        "如何解释当前全量测试未全绿？",
        "The audited snapshot collected 171 tests, with 161 passing and 10 failing. The failures concentrate in challenge thresholds, proof-attack state semantics, release-pack status, and an MCP backend-field contract. I do not hide them or describe the full suite as green. They reveal schema and expectation drift that should be resolved before strengthening the research claim.",
        None,
    ),
    (
        "How is Repo Agent different from related systems?",
        "Repo Agent 与相关工作有什么区别？",
        "GraphCodeBERT learns code representations with data-flow awareness. RepoCoder focuses on repository-level retrieval for code completion. Agentless separates localization, repair, and validation for real issues. LocAgent uses heterogeneous code graphs and LLM-guided localization. Repo Agent uses lightweight deterministic retrieval as its core and emphasizes route-grounded, replayable evidence, decoy audits, and change-governance artifacts. These are design differences, not proof of superior performance.",
        None,
    ),
    (
        "What is the most defensible innovation claim?",
        "最经得起质疑的创新点是什么？",
        "The strongest claim is not a new BM25 or PageRank algorithm. It is an evidence-oriented repository-localization loop that combines independent structural views, route-aware graph ranking, explicit hard-negative auditing, and replayable evidence contracts across code changes. The claim remains conditional: its external and downstream benefits must be demonstrated with stronger controlled experiments.",
        None,
    ),
    (
        "What are the biggest limitations, and what if external experiments show no gain?",
        "最大局限是什么？如果外部实验没有提升怎么办？",
        "The biggest limitations are incomplete parser coverage, heuristic graph resolution and reranking, uncalibrated confidence, incomplete external baselines, current full-suite failures, and no fixed-budget downstream repair study. If stronger external experiments show no gain, I would report the negative result, identify which components fail to generalize, retain any useful artifact or diagnostic contribution, and simplify or redirect the method rather than tune on the test set.",
        None,
    ),
    (
        "What is your concrete research plan for the next month and the next three years?",
        "未来一个月和三年的具体计划是什么？",
        "In the next month, I would repair schema and test drift, freeze a clean baseline, complete matched external baselines, measure parser recall, and add calibration and downstream trial protocols. Over three years, I would combine static graphs with runtime traces, tests, and version history; study confidence and abstention; and evaluate whether evidence-aware agents improve real repair quality under fixed budgets. The long-term goal is a trustworthy software-engineering agent that knows when not to act.",
        None,
    ),
]


SECTIONS = [
    ("6.2 项目定位与端到端数据流", PITCH_AND_FLOW),
    ("6.3 数据模型、配置、Ignore 与缓存", MODELS_CONFIG_CACHE),
    ("6.4 Parser、Chunk 与仓库图", PARSING_GRAPH),
    ("6.5 BM25、RRF、Query Plan、PPR 与重排", RETRIEVAL_RANKING),
    ("6.6 Agent、Tools、Runtime、Bundle 与交互层", AGENT_RUNTIME_INTERFACES),
    ("6.7 Proof、Impact、Contract、Temporal 与 Court", PROOF_GOVERNANCE),
    ("6.8 Engineering Mode、安全、CLI、Web 与 MCP", ENGINEERING_SECURITY),
    ("6.9 评测、可信度、Related Work、局限与未来", EVALUATION_RESEARCH),
]


THREE_MINUTE = [
    "My project is called Repo Agent. It addresses a problem that appears before code generation: a coding agent must first identify the correct repository context. In a real codebase, the same business words may appear in public, admin, legacy, mock, test, and documentation code. If the agent chooses the wrong implementation, even a well-written patch can be harmful.",
    "The system parses Python with AST and JavaScript or TypeScript mainly with Tree-sitter, then creates symbol-aware chunks and a typed repository graph. It builds independent BM25 views for content, identifiers, paths, and structure, combines their rankings with weighted reciprocal rank fusion, and uses query planning, exact route anchors, contrastive rules, and bounded Personalized PageRank to distinguish the relevant execution path from hard negatives.",
    "The result is not only an answer. It includes ranked files and symbols, line ranges, snippets, ranking reasons, graph paths, confidence warnings, a decoy audit, and a machine-readable proof object. Strict replay can later verify whether the saved nodes and graph edges still hold after the repository changes. The same evidence can support impact analysis, regression contracts, PR guards, and temporal drift diagnosis.",
    "The ten-case bundled suite reaches perfect Top-1, but I treat it as a regression signal. The thirty-two-case challenge reaches 84.375 percent Top-1, and a repository-disjoint external test shows smaller, statistically inconclusive improvements. Therefore, my current claim is an engineering feasibility result, not external state of the art. My next work is parser-recall evaluation, confidence calibration, stronger baselines, and a fixed-budget study of downstream repair quality.",
]


FIVE_MINUTE = [
    "Repo Agent starts from a simple observation: for repository-level tasks, generation quality depends on localization quality. A large language model can produce fluent code while reading the wrong file. Keyword search is excellent for exact strings, but it cannot by itself explain route-to-handler-to-writer relations or reject an admin function that shares the same words as a public endpoint.",
    "The repository is first validated and indexed. Python uses the standard AST. JavaScript and TypeScript prefer Tree-sitter, with a guarded fallback for large or unstable files. Each function, class, method, and route becomes a CodeChunk with its text, path, line range, calls, references, imports, inheritance, route, and parser provenance. The indexer also creates file summaries and a typed graph with call, reference, import, inheritance, and route-to-handler edges.",
    "At query time, a rule-based planner identifies intents such as API lookup, flow tracing, configuration, tests, front end, or styles. A file scout proposes likely modules, but a global chunk channel protects recall. Four independent BM25 views score code content, identifiers, paths, and structure. Weighted RRF combines their ranks because the raw score scales are not directly comparable. Route literals and contrastive exclusions provide precise evidence, while bounded PPR spreads relevance through a local graph neighborhood.",
    "The agent constructs a deterministic answer by default, so the system works without a model, API key, or GPU. Optional model reranking and tool calling operate only on bounded repository tools and observed evidence. The output contains hits, snippets, reasons, trace events, diagnostics, graph-search metadata, and a proof bundle. The proof is an engineering evidence contract, not a formal semantic proof.",
    "Proof replay checks whether the target, route, supporting path, and selected edges still exist. Impact analysis expands from the target to possible dependents. A regression contract freezes important evidence invariants, and a PR guard requires verification when changed files touch them. Temporal analysis can replay the contract across git commits, locate the first failure, infer a possible successor, and propose a reviewable JSON Patch migration.",
    "The evaluation deliberately separates claims. Bundled cases test regression; challenge and hard-negative cases test failure modes; repository-disjoint data tests generalization; mutation tests evaluate replay; and engineering artifacts test governance. Current external results are preliminary, the confidence score is not calibrated, and the full suite has known failures. I present those limits openly because the research goal is not an agent that guesses more confidently, but one that can show evidence and abstain when evidence is weak.",
]


TEN_MINUTE = [
    "I will explain Repo Agent through the problem, representation, retrieval, graph reasoning, evidence, engineering, evaluation, and limitations. The problem is repository-level localization. A user asks where a behavior is implemented or which function should be inspected. The answer may require several relations, and similar names can create hard negatives.",
    "The system boundary comes first. The input is a validated repository path, a natural-language question, and top-k. The output is a ranked evidence object. The default path does not edit code, train a model, require a GPU, or promise complete program semantics. Optional engineering mode is separated from investigation and uses a workspace.",
    "For representation, Python uses AST, while JavaScript and TypeScript mainly use Tree-sitter. HTML and CSS provide lighter link and import facts. Every extracted symbol becomes a CodeChunk with text, line positions, symbol metadata, route information, relations, and parser provenance. The graph contains calls, references, imports, inheritance, and route-to-handler edges. This graph is approximate and its quality depends on parser and name-resolution recall.",
    "For retrieval, the project uses four BM25 views: content, identifier, path, and structure. BM25 provides term-frequency saturation and length normalization. The views are fused with weighted reciprocal rank fusion, which uses ranks instead of adding incompatible raw scores. A file scout reduces noise, while a global chunk channel protects recall when scouting is wrong.",
    "For question-specific ranking, a deterministic planner identifies roles, languages, routes, actions, and symbol types. Exact route anchors distinguish route families. Contrastive rules downrank admin, legacy, mock, test, or documentation candidates when they conflict with the query. Bounded Personalized PageRank diffuses relevance from retrieved seeds through a local weighted graph. The active graph algorithm is PPR, although some compatibility fields still contain historical MCTS names.",
    "For answer construction, deterministic mode produces a grounded result without a model. Optional model reranking behaves like a cross-encoder over retrieved candidates, and optional tool calling can plan, read files, search symbols, follow relations, and observe verification results. Tools are bounded by deterministic path and command policy. The model cannot grant itself new permissions through repository text.",
    "The evidence layer stores the top hit, route anchors, supporting paths, proof graph, snippets, score gaps, warnings, and decoys. Strict replay verifies selected graph edges after repository change. Mutation tests intentionally corrupt evidence to check whether replay detects it. The proof is not a mathematical proof; it is a replayable engineering contract over the indexed evidence.",
    "The same artifact supports governance. Impact analysis identifies exposed routes and related files. Regression contracts freeze evidence invariants. PR guards request replay when a protected surface changes and can emit SARIF. Temporal analysis checks contracts across git history, identifies the first failing commit, explains graph deltas, ranks possible successor symbols, and proposes a human-reviewed migration plan. The evidence court adds role-specific claims and challenges rather than relying on agent voting.",
    "Evaluation is layered. The ten-case portable suite is a regression gate and reaches 100 percent Top-1. The thirty-two-case challenge reaches 84.375 percent Top-1 and exposes configuration, front-end state, and cross-file weaknesses. A two-hundred-query, twenty-two-repository collection uses repository-disjoint splits and a fifty-query frozen test. Full Multiview improves some top-k metrics over BM25, but the current paired bootstrap intervals cross zero. The full audited test run also has ten known failures concentrated in advanced artifact contracts.",
    "My conclusion is therefore limited but useful. The project demonstrates a complete evidence-oriented localization loop and exposes where it fails. The strongest future work is not adding more interface features. It is measuring parser recall, simplifying rule interactions, calibrating confidence and abstention, completing fair external baselines, and testing whether proof bundles improve real repair success under the same model and token budget.",
]


def update_existing_text(doc):
    replacements = {
        "Repo Agent 保研英语面试全题库": "Repo Agent 保研英语面试与完整项目答辩手册",
        "108 个高频问题 · 项目定制英文答案 · 连续追问与救场话术": "233 个高频问题 · 125 个项目细节追问 · 30 秒至 10 分钟口述稿",
        "Repo Agent 连续技术追问：32 题，覆盖 BM25、RRF、PPR、route anchor、proof、parser、评测、局限与个人贡献。": "Repo Agent 连续技术追问：32 题；新增项目全链路口述追问 125 题，覆盖全部核心模块、扩展系统、实验、源码细节与研究边界。",
        "附录：联网检索来源与项目内依据": "第六章：完整项目英语口述答辩；附录：联网检索来源与项目内依据",
    }
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                for run in paragraph.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        break
                else:
                    paragraph.text = paragraph.text.replace(old, new)
    doc.core_properties.title = "Repo Agent 保研英语面试与完整项目答辩手册"
    doc.core_properties.subject = "计算机保研英语口语、Repo Agent 全链路项目介绍与源码答辩"


def add_extension(doc):
    doc.add_heading("第六章  完整项目英语口述答辩", level=1)
    lead = doc.add_paragraph(style="Answer")
    label = lead.add_run("How to use this chapter  ")
    set_run_font(label, color=BLUE, bold=True)
    run = lead.add_run(
        "Start with the three-minute version. Then learn one category at a time and practice continuous follow-ups. The answers are written as spoken English: direct conclusion first, technical evidence second, and limitation last. Do not memorize every sentence mechanically; preserve the facts and logic."
    )
    set_run_font(run)

    doc.add_heading("6.1 多时长项目口述稿", level=2)
    add_script(doc, "6.1.1 三分钟标准项目介绍", THREE_MINUTE)
    add_script(doc, "6.1.2 五分钟完整项目介绍", FIVE_MINUTE)
    add_script(doc, "6.1.3 十分钟深度答辩稿", TEN_MINUTE)

    number = 109
    for heading, questions in SECTIONS:
        doc.add_heading(heading, level=2)
        for en, zh, answer, follow_up in questions:
            add_question(doc, number, en, zh, answer, follow_up)
            number += 1
    assert number == 234, number


def move_new_elements_before_anchor(doc, anchor, original_ids):
    body = doc.element.body
    new_elements = [el for el in list(body) if id(el) not in original_ids and el.tag != qn("w:sectPr")]
    for element in new_elements:
        anchor.addprevious(element)


def audit(doc):
    questions = []
    for p in doc.paragraphs:
        match = re.match(r"^Q(\d+)\.", p.text)
        if match:
            questions.append(int(match.group(1)))
    assert questions == list(range(1, 234)), (len(questions), questions[:3], questions[-3:])
    answers = sum(1 for p in doc.paragraphs if p.text.startswith("Suggested answer") or p.text.startswith("Oral answer"))
    assert answers == 233, answers
    assert any("十分钟深度答辩稿" in p.text for p in doc.paragraphs)
    assert any("Q233." in p.text for p in doc.paragraphs)
    assert any("233 个高频问题" in p.text for p in doc.paragraphs)
    assert len(PITCH_AND_FLOW) == 10
    assert len(MODELS_CONFIG_CACHE) == 10
    assert len(PARSING_GRAPH) == 16
    assert len(RETRIEVAL_RANKING) == 18
    assert len(AGENT_RUNTIME_INTERFACES) == 16
    assert len(PROOF_GOVERNANCE) == 20
    assert len(ENGINEERING_SECURITY) == 16
    assert len(EVALUATION_RESEARCH) == 19


def main():
    doc = Document(DOCX)
    update_existing_text(doc)
    anchor_paragraph = next(p for p in doc.paragraphs if p.text.startswith("附录  联网检索来源与项目内依据"))
    anchor = anchor_paragraph._p
    original_ids = {id(el) for el in list(doc.element.body)}
    add_extension(doc)
    move_new_elements_before_anchor(doc, anchor, original_ids)
    audit(doc)
    doc.save(DOCX)
    print(f"Updated {DOCX}")
    print("Questions: 233")
    print("New project-detail questions: 125")


if __name__ == "__main__":
    main()
